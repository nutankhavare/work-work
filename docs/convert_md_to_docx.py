import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font to Times New Roman style (simulated)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()

        # Handle Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator lines like | --- | --- |
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            
            cells = [c.strip() for c in line.split('|') if c.strip() != '']
            table_data.append(cells)
            continue
        else:
            if in_table:
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for r, row_cells in enumerate(table_data):
                        for c, cell_text in enumerate(row_cells):
                            # Clean up bold/italic markers for tables
                            clean_text = cell_text.replace('**', '').replace('_', '').replace('*', '')
                            table.cell(r, c).text = clean_text
                in_table = False
                table_data = []

        # Handle Headers
        if line.startswith('#'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            # Clean up bold markers in headers
            text = text.replace('**', '')
            h = doc.add_heading(text, level=level)
            # Center Chapter headers
            if 'CHAPTER' in text.upper() and level == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Handle Horizontal Rules
        if line == '---':
            doc.add_page_break()
            continue

        # Handle List Items
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Basic bold parsing
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            continue
        
        if re.match(r'^\d+\.', line):
            text = line[line.find('.')+1:].strip()
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            continue

        # Handle Regular Paragraphs
        if line != '':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_formatted_text(p, line)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

def _add_formatted_text(paragraph, text):
    # Very basic bold parser
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    target_dir = r"c:\Users\lenovo\my work\institute furthur\docs"
    md_file = os.path.join(target_dir, "VanLoka_Project_Report_Final.md")
    docx_file = os.path.join(target_dir, "VanLoka_Project_Report_Final.docx")
    
    if os.path.exists(md_file):
        convert_md_to_docx(md_file, docx_file)
    else:
        print(f"Error: {md_file} not found.")
